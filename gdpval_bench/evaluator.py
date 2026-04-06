"""
LLM-based Work Evaluator using Category-Specific Meta-Prompts.

Self-contained evaluator for GDPVal benchmark — no external dependencies
on ClawWork / livebench.  Meta-prompt rubrics are bundled in
``gdpval_bench/meta_prompts/``.

Originally adapted from ClawWork ``livebench/work/llm_evaluator.py``.
"""

import os
import json
import base64
import logging
from typing import Dict, Optional, Tuple, List, Union
from pathlib import Path
from openai import OpenAI
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

logger = logging.getLogger(__name__)

# Default meta-prompts directory (bundled with this package)
_DEFAULT_META_PROMPTS = Path(__file__).resolve().parent / "meta_prompts"


class LLMEvaluator:
    """
    LLM-based evaluator that uses category-specific meta-prompts
    to evaluate agent work artifacts with a 0.0-1.0 score.
    """

    def __init__(
        self,
        meta_prompts_dir: str = "",
        model: str = "gpt-4o",
    ):
        if not meta_prompts_dir:
            meta_prompts_dir = str(_DEFAULT_META_PROMPTS)
        self.meta_prompts_dir = Path(meta_prompts_dir)
        self.model = model

        api_key = os.getenv("EVALUATION_API_KEY") or os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError("Neither EVALUATION_API_KEY nor OPENAI_API_KEY found in environment")

        base_url = os.getenv("EVALUATION_API_BASE") or os.getenv("OPENAI_API_BASE")

        if os.getenv("EVALUATION_MODEL"):
            self.model = os.getenv("EVALUATION_MODEL")

        if os.getenv("EVALUATION_API_KEY"):
            print("🔧 Evaluation using separate API key (EVALUATION_API_KEY)")
        else:
            print("🔧 Evaluation using shared API key (OPENAI_API_KEY)")

        if base_url:
            print(f"🔧 Evaluation API base URL: {base_url}")
            self.client = OpenAI(api_key=api_key, base_url=base_url)
        else:
            print("🔧 Evaluation using default OpenAI endpoint")
            self.client = OpenAI(api_key=api_key)

        print(f"🔧 Evaluation model: {self.model}")

        self._meta_prompt_cache: Dict[str, Dict] = {}

    # ── Public API ────────────────────────────────────────────────

    def evaluate_artifact(
        self,
        task: Dict,
        artifact_paths: list[str],
        description: str = "",
        max_payment: float = 50.0,
    ) -> Tuple[float, str, float]:
        """Evaluate work artifact(s).

        Returns (evaluation_score 0.0-1.0, feedback_text, payment_amount).
        """
        occupation = task.get('occupation', '')
        if not occupation:
            return (0.0, "Error: Task missing occupation field", 0.0)

        meta_prompt = self._load_meta_prompt(occupation)
        if not meta_prompt:
            raise FileNotFoundError(
                f"No meta-prompt found for occupation '{occupation}'. "
                f"Check that meta_prompts/ contains the appropriate file."
            )

        # Read artifacts directly; let _read_artifacts_with_images raise on
        # missing/unreadable files rather than pre-checking existence (TOCTOU).
        artifact_data = self._read_artifacts_with_images(artifact_paths)
        user_message_content = self._build_multimodal_evaluation_content(
            meta_prompt=meta_prompt,
            task=task,
            artifact_data=artifact_data,
            missing_artifacts=[],
            description=description,
        )

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system",
                     "content": "You are an expert work evaluator. Follow the provided rubric precisely and output a structured evaluation."},
                    {"role": "user", "content": user_message_content}
                ],
            )
            evaluation_text = response.choices[0].message.content
            score = self._extract_score(evaluation_text)
            normalized_score = score / 10.0
            payment = normalized_score * max_payment
            return (normalized_score, evaluation_text, payment)
        except Exception as e:
            error_msg = f"LLM evaluation failed: {e}"
            print(f"  {error_msg}")
            logger.error(error_msg, exc_info=True)
            raise RuntimeError(error_msg) from e

    # ── Meta-prompt loading ───────────────────────────────────────

    def _load_meta_prompt(self, occupation: str) -> Optional[Dict]:
        normalized = occupation.replace(' ', '_').replace(',', '')
        if normalized in self._meta_prompt_cache:
            return self._meta_prompt_cache[normalized]
        meta_prompt_path = self.meta_prompts_dir / f"{normalized}.json"
        if not meta_prompt_path.exists():
            print(f"⚠️ No meta-prompt found for occupation: {occupation}")
            print(f"   Looking for: {meta_prompt_path}")
            return None
        try:
            with open(meta_prompt_path, 'r', encoding='utf-8') as f:
                meta_prompt = json.load(f)
            self._meta_prompt_cache[normalized] = meta_prompt
            return meta_prompt
        except Exception as e:
            print(f"⚠️ Error loading meta-prompt for {occupation}: {e}")
            return None

    # ── Artifact reading ──────────────────────────────────────────

    def _read_artifacts_with_images(
        self, artifact_paths: list[str], max_size_kb: int = 2000
    ) -> Dict[str, Dict[str, Union[str, bytes]]]:
        artifacts = {}
        for path in artifact_paths:
            file_size = os.path.getsize(path)
            file_ext = os.path.splitext(path)[1].lower()

            if file_size > max_size_kb * 1024:
                raise RuntimeError(f"File too large: {file_size} bytes (>{max_size_kb}KB) - {path}")
            if file_size == 0:
                raise ValueError(f"Empty file submitted for evaluation: {path}")

            if file_ext in ['.png', '.jpg', '.jpeg', '.gif', '.webp']:
                with open(path, 'rb') as f:
                    image_data = f.read()
                artifacts[path] = {'type': 'image', 'format': file_ext[1:], 'data': image_data, 'size': file_size}
            elif file_ext == '.docx':
                artifacts[path] = {'type': 'text', 'content': self._read_docx_content(path)}
            elif file_ext == '.xlsx':
                artifacts[path] = {'type': 'text', 'content': self._read_xlsx_content(path)}
            elif file_ext == '.pptx':
                artifacts[path] = {'type': 'text', 'content': f"[PPTX file: {file_size} bytes]"}
            elif file_ext == '.pdf':
                artifacts[path] = {'type': 'text', 'content': f"[PDF file: {file_size} bytes]"}
            else:
                try:
                    with open(path, 'r', encoding='utf-8') as f:
                        artifacts[path] = {'type': 'text', 'content': f.read()}
                except UnicodeDecodeError:
                    raise RuntimeError(f"Unsupported binary file type: {file_ext} - {path}")
        return artifacts

    def _read_docx_content(self, path: str) -> str:
        try:
            from docx import Document
        except ImportError:
            return f"[DOCX file present but python-docx not installed]"
        try:
            doc = Document(path)
            content = [f"[DOCX Document - {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables]\n"]
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
            for i, table in enumerate(doc.tables):
                content.append(f"\n--- Table {i+1} ({len(table.rows)} rows x {len(table.columns)} cols) ---")
                for row in table.rows[:10]:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        content.append(row_text)
                if len(table.rows) > 10:
                    content.append(f"... ({len(table.rows) - 10} more rows)")
            return "\n".join(content)
        except Exception as e:
            return f"[DOCX extraction failed: {e}]"

    def _read_xlsx_content(self, path: str) -> str:
        try:
            from openpyxl import load_workbook
        except ImportError:
            return f"[XLSX file present but openpyxl not installed]"
        try:
            wb = load_workbook(path, data_only=True)
            content = [f"[Excel Workbook - {len(wb.sheetnames)} sheets: {', '.join(wb.sheetnames)}]\n"]
            for sheet_name in wb.sheetnames[:5]:
                ws = wb[sheet_name]
                content.append(f"\n=== Sheet: {sheet_name} ({ws.max_row} rows x {ws.max_column} cols) ===")
                for row_idx, row in enumerate(ws.iter_rows(max_row=20, values_only=True), 1):
                    row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    if row_text.strip():
                        content.append(f"Row {row_idx}: {row_text}")
                if ws.max_row > 20:
                    content.append(f"... ({ws.max_row - 20} more rows)")
            if len(wb.sheetnames) > 5:
                content.append(f"\n... ({len(wb.sheetnames) - 5} more sheets)")
            return "\n".join(content)
        except Exception as e:
            return f"[XLSX extraction failed: {e}]"

    # ── Evaluation prompt building ────────────────────────────────

    def _build_multimodal_evaluation_content(
        self,
        meta_prompt: Dict,
        task: Dict,
        artifact_data: Dict[str, Dict],
        missing_artifacts: list[str],
        description: str
    ) -> List[Dict[str, Union[str, Dict]]]:
        evaluation_prompt = meta_prompt.get('evaluation_prompt', '')
        text_content = f"""# TASK EVALUATION REQUEST

## Category: {meta_prompt.get('category', 'Unknown')}

## Evaluation Guidelines:
{evaluation_prompt}

## Task Prompt (Original Assignment):
{task.get('prompt', 'N/A')}

## Task Metadata:
- Task ID: {task.get('task_id', 'N/A')}
- Sector: {task.get('sector', 'N/A')}
- Occupation: {task.get('occupation', 'N/A')}
- Reference Files: {', '.join(task.get('reference_files', [])) or 'None'}

## Agent's Description:
{description or 'No description provided'}

## Submitted Artifacts:

"""
        for path, artifact in artifact_data.items():
            if artifact['type'] == 'text':
                text_content += f"\n### File: {os.path.basename(path)}\n```\n{artifact['content']}\n```\n\n"
            elif artifact['type'] == 'image':
                text_content += f"\n### Image: {os.path.basename(path)} ({artifact['format']}, {artifact['size']} bytes)\n[See image below]\n\n"

        if missing_artifacts:
            text_content += "\n## Missing Artifacts:\n"
            for path in missing_artifacts:
                text_content += f"- {path}\n"

        text_content += """

---

Please evaluate this work according to the rubric above. Output your evaluation in this format:

**OVERALL SCORE:** [0-10]

**DIMENSION SCORES:**
[List dimension scores from rubric]

**KEY FINDINGS:**
[2-3 bullet points on what worked / didn't work]

**FEEDBACK:**
[1-2 paragraph explanation]

**TOP IMPROVEMENTS NEEDED:**
[Numbered list of 3 specific improvements]
"""
        content: List[Dict[str, Union[str, Dict]]] = [{"type": "text", "text": text_content}]

        for path, artifact in artifact_data.items():
            if artifact['type'] == 'image':
                image_base64 = base64.b64encode(artifact['data']).decode('utf-8')
                format_to_mime = {'png': 'image/png', 'jpg': 'image/jpeg', 'jpeg': 'image/jpeg',
                                  'gif': 'image/gif', 'webp': 'image/webp'}
                mime_type = format_to_mime.get(artifact['format'], 'image/png')
                content.append({
                    "type": "image_url",
                    "image_url": {"url": f"data:{mime_type};base64,{image_base64}", "detail": "high"}
                })
        return content

    # ── Score extraction ──────────────────────────────────────────

    def _extract_score(self, evaluation_text: str) -> float:
        import re
        patterns = [
            r'OVERALL SCORE:\s*(\d+(?:\.\d+)?)',
            r'Overall Score:\s*(\d+(?:\.\d+)?)',
            r'Score:\s*(\d+(?:\.\d+)?)/10',
            r'Final Score:\s*(\d+(?:\.\d+)?)',
        ]
        for pattern in patterns:
            match = re.search(pattern, evaluation_text, re.IGNORECASE)
            if match:
                return max(0.0, min(10.0, float(match.group(1))))
        first_part = evaluation_text[:200]
        numbers = re.findall(r'\b(\d+(?:\.\d+)?)\b', first_part)
        if numbers:
            score = float(numbers[0])
            if 0 <= score <= 10:
                return score
        print("⚠️ Could not extract score from evaluation, defaulting to 5.0")
        return 5.0
